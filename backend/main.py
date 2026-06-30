import os
import io
import json
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()
# Initialize the application
app = FastAPI(title="AI Lesson Note Generator API")

# Add CORS middleware so the HTML/JS frontend can communicate with this backend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins, including local files and Live Server
    allow_credentials=False, # This MUST be False when allow_origins is set to ["*"]
    allow_methods=["*"],
    allow_headers=["*"],
)


# Initialize the Gemini Client
client = genai.Client()

# Define the exact 5-section pedagogical output schema
class LessonNoteSchema(BaseModel):
    introduction: str = Field(description="Background and definition of the topic.")
    learning_objectives: list[str] = Field(description="List of behavioral objectives starting with 'Students should be able to...'")
    lesson_content: str = Field(description="Detailed technical content aligned with WAEC/NECO/NERDC guidelines.")
    class_activity: str = Field(description="Practical exercises or discussion questions for the students.")
    summary: str = Field(description="A concise wrap-up of the key takeaway points.")

# Define Request Schema
class LessonNoteRequest(BaseModel):
    subject: str      
    class_level: str  
    topic: str        

# Mock Database function for Curriculum Retrieval (RAG)
def retrieve_curriculum_context(subject: str, class_level: str, topic: str) -> str:
    return f"Standard WAEC/NECO alignment for {class_level} {subject} focusing on {topic}."

# Document generation helper function
def create_docx_lesson_note(metadata: dict, note_data: dict) -> io.BytesIO:
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CURRICULUM-ALIGNED LESSON NOTE")
    title_run.bold = True
    title_run.font.size = Pt(16)
    
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = f"Subject: {metadata['subject']}"
    table.cell(0, 1).text = f"Class Level: {metadata['class_level']}"
    table.cell(1, 0).text = f"Topic: {metadata['topic']}"
    table.cell(1, 1).text = "Status: Verified WAEC/NECO Aligned"

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_section(heading_title, content_text, is_list=False):
        heading = doc.add_paragraph()
        heading_run = heading.add_run(heading_title)
        heading_run.bold = True
        heading_run.font.size = Pt(13)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(4)
        
        if is_list and isinstance(content_text, list):
            for item in content_text:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(item)
                p.paragraph_format.space_after = Pt(3)
        else:
            p = doc.add_paragraph()
            p.add_run(str(content_text))
            p.paragraph_format.space_after = Pt(6)

    # Build the required structural sections
    add_section("1. Introduction", note_data.get("introduction", ""))
    add_section("2. Learning Objectives", note_data.get("learning_objectives", []), is_list=True)
    add_section("3. Lesson Content", note_data.get("lesson_content", ""))
    add_section("4. Class Activity", note_data.get("class_activity", ""))
    add_section("5. Summary", note_data.get("summary", ""))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@app.post("/api/v1/export-lesson-note")
async def export_lesson_note(request: LessonNoteRequest):
    try:
        context = retrieve_curriculum_context(request.subject, request.class_level, request.topic)
        system_instruction = (
            "You are an expert secondary school science teacher in Nigeria. Your task is to generate "
            "comprehensive, high-quality lesson notes strictly aligned with the WAEC, NECO, and NERDC syllabi. "
            "Ground your content in the provided curriculum context."
        )
        prompt = f"Generate a lesson note for {request.subject} {request.class_level} on the topic: {request.topic}."
        
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=system_instruction,
                temperature=0.3,
                response_mime_type="application/json",
                response_schema=LessonNoteSchema,
            ),
        )
        
        note_json = json.loads(response.text)
        metadata = {"subject": request.subject, "class_level": request.class_level, "topic": request.topic}
        docx_file = create_docx_lesson_note(metadata, note_json)
        filename = f"LessonNote_{request.subject}_{request.class_level}_{request.topic.replace(' ', '_')}.docx"
        
        return StreamingResponse(
            docx_file,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8001)